import streamlit as st
import numpy as np
import easyocr as ocr
import io
from PIL import Image, ImageOps
import pandas as pd
import cv2
from docx import Document

st.set_page_config(layout="wide")
st.markdown("#### Оптическое распознавание таблиц и текста")

@st.cache_resource()
def load_models(langs):
    try:
        reader = ocr.Reader(langs, model_storage_directory=".")
        return reader
    except Exception as e:
        st.error(f"Ошибка при загрузке моделей EasyOCR: {e}")
        return None

available_langs = ["ru", "en", "es", "fr", "de"]
default_langs = ["ru", "en"]
selected_langs = st.multiselect(
    "Выберите языки для распознавания:",
    available_langs,
    default=default_langs,
)

reader = load_models(selected_langs)
if reader is None:
    st.stop()

def resize_image(image):
    img_array = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

    desired_width = 800
    height, width = img_array.shape[:2]
    aspect_ratio = width / height
    new_height = int(desired_width / aspect_ratio)
    img_resized = cv2.resize(img_array, (desired_width, new_height))

    return img_resized

def detect_lines(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    edges = cv2.Canny(blurred, 50, 150, apertureSize=3)
    
    # Настроим параметры HoughLinesP для лучшего обнаружения линий таблиц
    lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=40, minLineLength=500, maxLineGap=280)
    
    return lines is not None

def find_tables(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

    # Горизонтальные линии
    horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
    remove_horizontal = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
    cnts = cv2.findContours(remove_horizontal, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cnts = cnts[0] if len(cnts) == 2 else cnts[1]
    for c in cnts:
        cv2.drawContours(thresh, [c], -1, (0, 0, 0), 5)

    # Вертикальные линии
    vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
    remove_vertical = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
    cnts = cv2.findContours(remove_vertical, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cnts = cnts[0] if len(cnts) == 2 else cnts[1]
    for c in cnts:
        cv2.drawContours(thresh, [c], -1, (0, 0, 0), 5)

    # Инвертируем изображение для поиска таблиц
    invert = 255 - thresh

    # Морфологическое закрытие для соединения близких элементов
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (10, 10))
    morphed = cv2.morphologyEx(invert, cv2.MORPH_CLOSE, kernel)

    # Находим контуры таблиц
    cnts = cv2.findContours(morphed, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cnts = cnts[0] if len(cnts) == 2 else cnts[1]

    table_bboxes = []
    for c in cnts:
        x, y, w, h = cv2.boundingRect(c)
        table_bboxes.append((x, y, w, h))
    
    return table_bboxes

def extract_table_data(results, table_bboxes):
    table_data = []
    text_data = []

    for bbox, text, prob in results:
        x1, y1 = int(bbox[0][0]), int(bbox[0][1])
        x2, y2 = int(bbox[2][0]), int(bbox[2][1])
        
        is_in_table = False
        for x, y, w, h in table_bboxes:
            if x1 >= x and y1 >= y and x2 <= x+w and y2 <= y+h:
                is_in_table = True
                break

        if is_in_table:
            # Добавляем текст в таблицу
            table_data.append(text)

        text_data.append(text)

            # Преобразуем плоский список table_data в двумерный список (строки таблицы)
    table_rows = []
    current_row = []
    row_y = -1  # Инициализируем значением, которое не может быть координатой
    row_height_threshold = 15  # Порог высоты строки
    
    for bbox, text, prob in results:
        x1, y1 = int(bbox[0][0]), int(bbox[0][1])
        
        is_in_table = False
        for x, y, w, h in table_bboxes:
            if x1 >= x and y1 >= y and x1 <= x+w and y1 <= y+h:
                is_in_table = True
                break
        
        if is_in_table:
            if row_y == -1 or abs(y1 - row_y) > row_height_threshold:
                # Начинаем новую строку
                if current_row:
                    table_rows.append(current_row)
                current_row = [text]
                row_y = y1
            else:
                # Добавляем в текущую строку
                current_row.append(text)
    
    if current_row:
        table_rows.append(current_row)
    df = pd.DataFrame(text_data, table_rows)
    return df
    


uploaded_files = st.file_uploader("Загрузите изображения для распознавания", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

if uploaded_files:
    for uploaded_file in uploaded_files:
        image = Image.open(uploaded_file)
        st.image(image, caption=f"Загруженное изображение: {uploaded_file.name}", use_column_width=True)

        img_resized = resize_image(image)
        image_array = cv2.cvtColor(img_resized, cv2.COLOR_BGR2RGB)

        has_lines = detect_lines(image_array)
        if has_lines:
            st.warning("Обнаружены линии на изображении. Это может повлиять на точность распознавания.")
        
        table_bboxes = find_tables(image_array)

        results = reader.readtext(image_array)

        if table_bboxes:
            st.success("Таблицы обнаружены. Извлечение данных...")
            table_data_df = extract_table_data(results, table_bboxes)
            st.dataframe(table_data_df)
        else:
            st.info("Таблицы не обнаружены на изображении.")

        st.markdown("---")  # Отделитель между изображениями




def save_to_txt(text_data):
    txt_buffer = io.StringIO()
    txt_buffer.write(" ".join(text_data))
    txt_buffer.seek(0)
    return txt_buffer.getvalue()

def save_to_docx(table_data, text_data):
    # Создаем новый документ
    doc = Document()

    # Добавляем таблицу для распознанных данных
    if table_data:
        # Определяем количество столбцов
        num_columns = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=num_columns)

        for i, row in enumerate(table_data):
            for j, cell_value in enumerate(row):
                table.cell(i, j).text = str(cell_value)

    # Добавляем текст ниже таблицы
    if text_data:
        full_text = ' '.join(text_data)
        doc.add_paragraph(full_text)

    # Сохраняем документ в буфер
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def image_to_table(img_file_buffer):
    if img_file_buffer is not None:
        try:
            image = Image.open(img_file_buffer)
            image = ImageOps.exif_transpose(image)
            image_resized = resize_image(image)

            with st.expander("Изображение загружено"):
                st.image(image_resized, use_container_width=True)
            img_array = cv2.cvtColor(np.array(image_resized), cv2.COLOR_RGB2BGR)
            has_lines = detect_lines(img_array)
            table_bboxes = find_tables(img_array)

            with st.spinner("Распознавание..."):
                results = reader.readtext(image_resized, paragraph=False)
                table_data, text_data = extract_table_data(results, table_bboxes)

                return table_data, text_data
                
        except Exception as e:
            st.error(f"Ошибка при распознавании: {e}")
            return None, None
    return None, None

img_file_buffer = None
image_input = st.radio(
    "Выберите способ ввода текста:",
    ["Изображение", "Камера"],
    horizontal=True,
)

if image_input == "Камера":
    img_file_buffer = st.camera_input("Сделайте фото", key="camera_input")
elif image_input == "Изображение":
    img_file_buffer = st.file_uploader(
        "Загрузите изображение", type=["png", "jpg", "jpeg"]
    )

if img_file_buffer:
    extracted_data, extracted_text = image_to_table(img_file_buffer)
    
    if extracted_text:
        st.markdown("##### Распознанный текст")
        st.write(" ".join(extracted_text))

        # Кнопка для скачивания текста в формате TXT
        txt_data = save_to_txt(extracted_text)
        st.download_button(
            label="Скачать TXT",
            data=txt_data,
            file_name="extracted_text.txt",
            mime="text/plain",
        )

        # Кнопка для скачивания текста в формате DOCX
        docx_buffer = save_to_docx(extracted_data, extracted_text)
        st.download_button(
            label="Скачать DOCX",
            data=docx_buffer,
            file_name="extracted_text.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if extracted_data is not None:
        st.markdown("##### Распознанная таблица")
       
        df = pd.DataFrame(extracted_data)
        st.data_editor(extracted_data)
        xlsx_buffer = io.BytesIO()
        with pd.ExcelWriter(xlsx_buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, header=False)

        xlsx_buffer.seek(0)
        st.download_button(
            label="Скачать XLSX",            
            data=xlsx_buffer,
            file_name="extracted_table.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    
