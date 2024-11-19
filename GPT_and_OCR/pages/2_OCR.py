import streamlit as st
import requests
from PIL import Image
from io import BytesIO
import pytesseract
from docx import Document
import cv2
import numpy as np
import pdfplumber


def preprocess_image(image, blur_value, threshold_value):
    """Функция для предварительной обработки изображения."""
    gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
    gray = cv2.medianBlur(gray, blur_value)
    _, thresh = cv2.threshold(gray, threshold_value, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return Image.fromarray(thresh)

st.title("Оцифровка изображения")

# Sidebar parameters
blur_value = st.sidebar.slider("Выберите уровень размытия", min_value=1, max_value=15, value=3, step=2)
threshold_value = st.sidebar.slider("Выберите порог для обработки", min_value=0, max_value=255, value=0)

# Выбор языка
lang = st.sidebar.selectbox(
    "Выберите язык для распознавания текста",
    options=["rus+eng", "eng", "rus"],
    index=0
)

image_input = st.text_input("Вставьте ссылку на изображение")

def load_image_from_url(image_input):
    response = requests.get(image_input)
    response.raise_for_status()
    return Image.open(BytesIO(response.content))

if image_input:
    try:
        if "http" in image_input:
            image = load_image_from_url(image_input)
        else:
            image = Image.open(image_input)
    except Exception as e:
        st.error(f"Ошибка при загрузке изображения: {e}")
    else:
        image = preprocess_image(image, blur_value, threshold_value)  # Предварительная обработка
        st.image(image, caption="Загруженное изображение", use_column_width=True)

        recognized_text = pytesseract.image_to_string(image, lang=lang, config='--psm 6')
        st.write(recognized_text)

uploaded_file = st.file_uploader("Загрузите изображение или PDF", ["jpg", "jpeg", "png", "gif", "bmp", "pdf"])

def extract_text_and_tables(pdf, tables, full_text):
    for i, page in enumerate(pdf.pages):
        text = page.extract_text()
        if text:
            full_text += text + "\n"

        extracted_tables = page.extract_tables()
        for table in extracted_tables:
            tables.append(table)

        img = page.to_image()
        img_pil = img.original.convert("RGB")
        img_processed = preprocess_image(img_pil, blur_value, threshold_value)  # Предварительная обработка
        st.image(img_processed, caption=f"Страница {i + 1}", use_column_width=True)
        recognized_text = pytesseract.image_to_string(img_processed, lang=lang, config='--psm 6')
        st.write(recognized_text)

if uploaded_file is not None:
    full_text = ""
    tables = []
    if uploaded_file.type == "application/pdf":
        with pdfplumber.open(uploaded_file) as pdf:
            extract_text_and_tables(pdf, tables, full_text)

    else:
        image = Image.open(uploaded_file)
        image = preprocess_image(image, blur_value, threshold_value)  # Предварительная обработка
        st.image(image, caption="Загруженное изображение", use_column_width=True)

        recognized_text = pytesseract.image_to_string(image, lang=lang, config='--psm 6')
        st.write(recognized_text)

    if tables:
        st.text("Извлеченные таблицы")
        for i, table in enumerate(tables):
            st.text(f"Таблица {i + 1}")
            for row in table:
                st.write(" | ".join(str(cell) for cell in row))

    if st.button("Сохранить в TXT"):
        with open('recognized_text.txt', 'w', encoding='utf-8') as f:
            f.write(full_text if uploaded_file.type == "application/pdf" else recognized_text)

        st.download_button(
            label="Скачать TXT",
            data=open('recognized_text.txt', 'rb').read(),
            file_name='recognized_text.txt',
            mime='text/plain'
        )

    if st.button("Сохранить в DOCX"):
        doc = Document()
        doc.add_paragraph(full_text if uploaded_file.type == "application/pdf" else recognized_text)
        for i, table in enumerate(tables):
            doc.add_paragraph(f"Таблица {i + 1}")
            table_in_docx = doc.add_table(rows=len(table), cols=len(table[0]))
            for row_index, row in enumerate(table):
                cells = table_in_docx.rows[row_index].cells
                for cell_index, cell in enumerate(row):
                    cells[cell_index].text = str(cell)

        doc.save('recognized_text.docx')

        st.download_button(
            label="Скачать DOCX",
            data=open('recognized_text.docx', 'rb').read(),
            file_name='recognized_text.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )