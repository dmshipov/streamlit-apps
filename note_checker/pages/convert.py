import streamlit as st
import soundfile as sf
import speech_recognition as sr
import io
import os
import requests
from PIL import Image
import pytesseract
from io import BytesIO
from docx import Document
import cv2
import numpy as np
import pdfplumber
from moviepy.editor import VideoFileClip
from pydub import AudioSegment
import noisereduce as nr

def recognize_speech(audio_data, samplerate, language="ru-RU"):
    with io.BytesIO() as f:
        sf.write(f, audio_data, samplerate, format='WAV')
        f.seek(0)

        recognizer = sr.Recognizer()
        # Используем sr.AudioFile с правильным объектом
        with sr.AudioFile(f) as source:
            audio = recognizer.record(source)
            try:
                text = recognizer.recognize_google(audio, language=language)
                return text
            except sr.UnknownValueError:
                return "Речь не распознана"
            except sr.RequestError as e:
                return f"Ошибка сервиса распознавания речи: {e}"

def preprocess_audio(audio_file_path):
    audio = AudioSegment.from_wav(audio_file_path)
    samples = audio.get_array_of_samples()
    reduced_noise = nr.reduce_noise(y=samples, sr=audio.frame_rate)

    processed_audio = AudioSegment(
        reduced_noise.tobytes(),
        frame_rate=audio.frame_rate,
        sample_width=audio.sample_width,
        channels=audio.channels
    )

    processed_audio_path = audio_file_path.replace(".wav", "_processed.wav")
    processed_audio.export(processed_audio_path, format="wav")
    
    return processed_audio_path

def recognize_from_file(file, language="ru-RU"):
    recognizer = sr.Recognizer()

    try:
        with sr.AudioFile(file) as source:
            audio_length = source.DURATION
            step = 10  # шаг в секундах для извлечения фрагментов
            text = ""

            for start_time in range(0, int(audio_length), step):
                # Вместо source.seek(start_time) используем record с параметрами
                audio = recognizer.record(source, duration=step)

                try:
                    fragment_text = recognizer.recognize_google(audio, language=language)
                    text += fragment_text + " "
                except sr.UnknownValueError:
                    continue
                except sr.RequestError as e:
                    return f"Ошибка сервиса распознавания речи: {e}"

            return text.strip()

    except sr.UnknownValueError:
        return "Речь не распознана"
    except sr.RequestError as e:
        return f"Ошибка сервиса распознавания речи: {e}"
    except Exception as e:
        return f"Произошла ошибка: {e}"
    
def extract_audio_from_video(video_file):
    audio_file = "extracted_audio.wav"
    with VideoFileClip(video_file) as video:
        video.audio.write_audiofile(audio_file)
    return audio_file
    
def history_reset_function():
    # Code to be executed when the reset button is clicked
    st.session_state.clear() 



file_type = st.sidebar.radio("Выберите функцию:", ("Конвертация аудио в текст", "Конвертация изображения в текст"))

if file_type == "Конвертация аудио в текст":
    st.markdown("## Конвертация аудио в текст")
    uploaded_file = st.file_uploader("Загрузите файл MP3, WAV или видео для преобразования в текст", type=["mp3", "wav", "mp4", "avi", "mov"])
    
    if uploaded_file is not None:
        with open("temp_file", "wb") as f:
            f.write(uploaded_file.getbuffer())

        if uploaded_file.name.endswith((".mp3", ".wav")):
            recognized_text = recognize_from_file("temp_file")
        elif uploaded_file.name.endswith((".mp4", ".avi", ".mov")):
            audio_file = extract_audio_from_video("temp_file")
            recognized_text = recognize_from_file(audio_file)
            os.remove(audio_file)
        else:
            st.write("Пожалуйста, загрузите файл MP3, WAV или видео.")
            os.remove("temp_file")  # Удаляем временный файл
           
        st.write("Распознанный текст:")
        st.text(recognized_text)

        txt_file_name = "recognized_text.txt"
        st.download_button(
            label="Скачать распознанный текст (TXT)",
            data=recognized_text,
            file_name=txt_file_name,
            mime="text/plain"
        )

        os.remove("temp_file")


if file_type == "Конвертация изображения в текст":
    st.markdown("## Конвертация изображения в текст")
    def preprocess_image(image, blur_value, threshold_value):
        """Функция для предварительной обработки изображения."""
        gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
        gray = cv2.medianBlur(gray, blur_value)
        _, thresh = cv2.threshold(gray, threshold_value, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return Image.fromarray(thresh)

    

    # Sidebar parameters
    blur_value = st.sidebar.slider("Выберите уровень размытия", min_value=1, max_value=15, value=3, step=2)
    threshold_value = st.sidebar.slider("Выберите порог для обработки", min_value=0, max_value=255, value=100)

    # Выбор языка
    lang = st.sidebar.selectbox(
        "Выберите язык для распознавания текста",
        options=["rus+eng", "eng", "rus"],
        index=0
    )

    image_input = st.text_input("Вставьте ссылку для оцифровки текста")

    def load_image_from_url(image_input):
        response = requests.get(image_input)
        response.raise_for_status()
        return Image.open(BytesIO(response.content))

    if image_input:
        try:
            if "http" in image_input:
                image = load_image_from_url(image_input)
            else:
                image = Image.open(image_input)
        except Exception as e:
            st.error(f"Ошибка при загрузке изображения: {e}")
        else:
            image = preprocess_image(image, blur_value, threshold_value)  # Предварительная обработка
            st.image(image, caption="Загруженное изображение", use_column_width=True)

            recognized_text = pytesseract.image_to_string(image, lang=lang, config='--psm 6')
            st.write(recognized_text)

    uploaded_file = st.file_uploader("Загрузите изображение или PDF для оцифровки текста", ["jpg", "jpeg", "png", "gif", "bmp", "pdf"])

    def extract_text_and_tables(pdf, tables, full_text):
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                full_text += text + "\n"

            extracted_tables = page.extract_tables()
            for table in extracted_tables:
                tables.append(table)

            img = page.to_image()
            img_pil = img.original.convert("RGB")
            img_processed = preprocess_image(img_pil, blur_value, threshold_value)  # Предварительная обработка
            st.image(img_processed, caption=f"Страница {i + 1}", use_column_width=True)
            recognized_text = pytesseract.image_to_string(img_processed, lang=lang, config='--psm 6')
            st.write(recognized_text)

    if uploaded_file is not None:
        full_text = ""
        tables = []
        if uploaded_file.type == "application/pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                extract_text_and_tables(pdf, tables, full_text)

        else:
            image = Image.open(uploaded_file)
            image = preprocess_image(image, blur_value, threshold_value)  # Предварительная обработка
            st.image(image, caption="Загруженное изображение", use_column_width=True)

            recognized_text = pytesseract.image_to_string(image, lang=lang, config='--psm 6')
            st.write(recognized_text)

        if tables:
            st.text("Извлеченные таблицы")
            for i, table in enumerate(tables):
                st.text(f"Таблица {i + 1}")
                for row in table:
                    st.write(" | ".join(str(cell) for cell in row))

        if st.button("Сохранить в TXT"):
            with open('recognized_text.txt', 'w', encoding='utf-8') as f:
                f.write(full_text if uploaded_file.type == "application/pdf" else recognized_text)

            st.download_button(
                label="Скачать TXT",
                data=open('recognized_text.txt', 'rb').read(),
                file_name='recognized_text.txt',
                mime='text/plain'
            )

        if st.button("Сохранить в DOCX"):
            doc = Document()
            doc.add_paragraph(full_text if uploaded_file.type == "application/pdf" else recognized_text)
            for i, table in enumerate(tables):
                doc.add_paragraph(f"Таблица {i + 1}")
                table_in_docx = doc.add_table(rows=len(table), cols=len(table[0]))
                for row_index, row in enumerate(table):
                    cells = table_in_docx.rows[row_index].cells
                    for cell_index, cell in enumerate(row):
                        cells[cell_index].text = str(cell)

            doc.save('recognized_text.docx')

            st.download_button(
                label="Скачать DOCX",
                data=open('recognized_text.docx', 'rb').read(),
                file_name='recognized_text.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )